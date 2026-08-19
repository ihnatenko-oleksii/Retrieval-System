import json

from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.ingestion.loaders import CSVLoader, DocxLoader, JSONLoader, PDFLoader, TextLoader, get_loader


class TestTextLoader:
    def test_reads_utf8_content(self, tmp_path):
        path = tmp_path / "note.md"
        path.write_text("# Title\n\ncontent", encoding="utf-8")

        docs = TextLoader().load(str(path))

        assert len(docs) == 1
        assert docs[0].content == "# Title\n\ncontent"
        assert docs[0].metadata.extension == ".md"

    def test_falls_back_to_latin1_on_decode_error(self, tmp_path):
        path = tmp_path / "legacy.txt"
        # 0xe9 is not valid standalone UTF-8 but decodes cleanly as latin-1 ('é').
        path.write_bytes(b"Caf\xe9 note")

        docs = TextLoader().load(str(path))

        assert len(docs) == 1
        assert docs[0].content == "Café note"


class TestCSVLoader:
    def test_renders_rows_as_text(self, tmp_path):
        path = tmp_path / "data.csv"
        path.write_text("name,age\nAlice,30\nBob,25\n", encoding="utf-8")

        docs = CSVLoader().load(str(path))

        assert len(docs) == 1
        assert "Alice" in docs[0].content
        assert "Bob" in docs[0].content


class TestJSONLoader:
    def test_pretty_prints_json_content(self, tmp_path):
        path = tmp_path / "data.json"
        path.write_text(json.dumps({"key": "value", "nested": {"a": 1}}), encoding="utf-8")

        docs = JSONLoader().load(str(path))

        assert len(docs) == 1
        assert '"key": "value"' in docs[0].content


class TestDocxLoader:
    def test_extracts_non_empty_paragraphs(self, tmp_path):
        path = tmp_path / "doc.docx"
        doc = DocxDocument()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("")  # blank paragraphs are dropped
        doc.add_paragraph("Second paragraph.")
        doc.save(str(path))

        docs = DocxLoader().load(str(path))

        assert len(docs) == 1
        assert "First paragraph." in docs[0].content
        assert "Second paragraph." in docs[0].content

    def test_document_with_no_text_yields_no_documents(self, tmp_path):
        path = tmp_path / "empty.docx"
        doc = DocxDocument()
        doc.save(str(path))

        docs = DocxLoader().load(str(path))

        assert docs == []


class TestPDFLoader:
    def test_blank_page_yields_no_documents(self, tmp_path):
        path = tmp_path / "blank.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        with open(path, "wb") as f:
            writer.write(f)

        docs = PDFLoader().load(str(path))

        assert docs == []


class TestGetLoader:
    def test_known_text_extensions_use_text_loader(self):
        for ext in [".txt", ".md", ".py", ".java", ".js", ".ts", ".yml", ".yaml", ".xml", ".properties"]:
            assert isinstance(get_loader(f"file{ext}"), TextLoader)

    def test_dispatches_by_extension(self):
        assert isinstance(get_loader("a.csv"), CSVLoader)
        assert isinstance(get_loader("a.json"), JSONLoader)
        assert isinstance(get_loader("a.docx"), DocxLoader)
        assert isinstance(get_loader("a.pdf"), PDFLoader)

    def test_unsupported_extension_returns_none(self):
        assert get_loader("archive.zip") is None
        assert get_loader("no_extension") is None

    def test_extension_matching_is_case_insensitive(self):
        assert isinstance(get_loader("NOTES.MD"), TextLoader)
